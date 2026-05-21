from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create compact document (1-2 pages)
doc = Document()

# Set narrow margins for compactness
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title
title = doc.add_heading('Multimodal Bioactivity Prediction with IBM MoLFormer-XL', 0)
title_format = title.paragraph_format
title_format.space_after = Pt(4)
title_format.space_before = Pt(0)
for run in title.runs:
    run.font.size = Pt(13)
    run.font.bold = True

# 1. Model Architecture
doc.add_heading('1. Molecular Representation Encoding', 1)
p = doc.add_paragraph(
    'The FineTunedBERTaECFP model integrates three molecular representations: SMILES (via IBM MoLFormer-XL-both-10pct), '
    'ECFP fingerprints, and physicochemical descriptors.'
)

# SMILES Encoding
p = doc.add_paragraph()
p.add_run('SMILES Encoding: ').bold = True
p.add_run('For tokenized sequence x = [x₁, ..., xₙ], the transformer produces:')

eq = doc.add_paragraph('H = MoLFormer(x) ∈ ℝⁿˣ⁷⁶⁸,    h_SMILES = (1/3)(h_CLS + h_mean + h_max)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ECFP Projection
p = doc.add_paragraph()
p.add_run('ECFP Projection: ').bold = True
p.add_run('Two-layer MLP projects 1024-bit fingerprint:')

eq = doc.add_paragraph('h_ECFP = Dropout_p(GELU(LN(W₂ · Dropout₀.₅ₚ(GELU(LN(W₁·f_ECFP))))))    [1024→512→768]', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Descriptor Projection
p = doc.add_paragraph()
p.add_run('Descriptor Projection: ').bold = True
p.add_run('Three-layer MLP for 7 descriptors (MW, LogP, NumHDonors, TPSA, NumRotatableBonds, FractionCSP3, RingCount):')

eq = doc.add_paragraph('h_desc = MLP₃(d_normalized)    [7→256→512→768]', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 2. Cross-Modal Attention
doc.add_heading('2. Cross-Modal Attention & Fusion', 1)
p = doc.add_paragraph()
p.add_run('Multi-Head Attention: ').bold = True
p.add_run('H heads allow structural/physicochemical features to attend to SMILES:')

eq = doc.add_paragraph('MultiHead(Q, K, V) = Concat(head₁, ..., headₕ)W^O,    headᵢ = softmax(QWᵢ^Q(KWᵢ^K)ᵀ/√dₖ)VWᵢ^V', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_ECFP^attn = LN(h_ECFP + MultiHead(h_ECFP, h_SMILES, h_SMILES))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_desc^attn = LN(h_desc + MultiHead(h_desc, h_SMILES, h_SMILES))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Bilinear Fusion: ').bold = True
p.add_run('Captures multiplicative interactions (best performing):')

eq = doc.add_paragraph('h_fused = h_SMILES^T W_B (h_ECFP^attn + h_desc^attn),    W_B ∈ ℝ⁷⁶⁸ˣ⁷⁶⁸', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 3. Classification Head & Training
doc.add_heading('3. Classification Head & Training', 1)
p = doc.add_paragraph()
p.add_run('Progressive Reduction: ').bold = True
p.add_run('N-layer MLP with halved dimensions:')

eq = doc.add_paragraph('h⁽ⁱ⁾ = Dropout_pᵢ(GELU(LN(Wᵢh⁽ⁱ⁻¹⁾))),    ŷ = w_out^T h⁽ᴺ⁾', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Loss Function: ').bold = True
p.add_run('Weighted BCE for class imbalance:')

eq = doc.add_paragraph('ℒ = -(1/B)Σᵢ[w_pos·yᵢlog(σ(ŷᵢ)) + (1-yᵢ)log(1-σ(ŷᵢ))],    w_pos = N_neg/N_pos', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Optimization: ').bold = True
p.add_run('RMSprop with OneCycleLR scheduler:')

eq = doc.add_paragraph('α(t) = αₘₐₓ·((1-t/T)/(cos(π·t/T)+1))·0.5 + αₘᵢₙ,    αₘᵢₙ = αₘₐₓ/25', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 4. Hyperparameter Optimization
doc.add_heading('4. Hyperparameter Optimization', 1)
p = doc.add_paragraph('Optuna TPE sampling over 50 trials with MedianPruner:')

# Compact hyperparameter table
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

hdr = table.rows[0].cells
hdr[0].text = 'Hyperparameter'
hdr[1].text = 'Range/Options'
hdr[2].text = 'Hyperparameter'
hdr[3].text = 'Range/Options'

params_left = [
    ('unfreeze_layers', '[4, 12]'),
    ('num_heads', '{4,8,12,16}'),
    ('fusion_type', '{concat,gated,bilinear}'),
    ('dropout', '[0.1, 0.6]')
]

params_right = [
    ('learning_rate', '[5e-7, 1e-5]'),
    ('weight_decay', '[1e-5, 1e-3]'),
    ('batch_size', '{16,32,64}'),
    ('grad_accum', '{1,2,4,8}')
]

for i, ((p1, r1), (p2, r2)) in enumerate(zip(params_left, params_right), 1):
    row = table.rows[i].cells
    row[0].text = p1
    row[1].text = r1
    row[2].text = p2
    row[3].text = r2

# Set font size for all cells after text is added
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# 5. Results
doc.add_heading('5. Results', 1)

p = doc.add_paragraph()
p.add_run('Best Configuration: ').bold = True
p.add_run('5 unfrozen layers, 12 heads, bilinear fusion, RMSprop + OneCycleLR, dropout=0.30, batch=32×2 accum')

p = doc.add_paragraph()
p.add_run('Dataset: ').bold = True
p.add_run('6,341 molecular compounds with SMILES, ECFP (radius=2, 1024 bits), 7 descriptors')

# Results table
table = doc.add_table(rows=3, cols=7)
table.style = 'Light Grid Accent 1'

hdr = table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'AUROC'
hdr[2].text = 'Bal. Acc.'
hdr[3].text = 'Precision'
hdr[4].text = 'Recall'
hdr[5].text = 'F1'
hdr[6].text = 'MCC'

row1 = table.rows[1].cells
row1[0].text = 'Classification'
row1[1].text = '95.39%'
row1[2].text = '92.01%'
row1[3].text = '97.12%'
row1[4].text = '96.24%'
row1[5].text = '94.60%'
row1[6].text = '0.8281'

row2 = table.rows[2].cells
row2[0].text = 'Regression'
row2[1].text = '-'
row2[2].text = '-'
row2[3].text = 'R²: 0.7006'
row2[4].text = 'RMSE: 0.7055'
row2[5].text = 'MAE: 0.4932'
row2[6].text = 'r: 0.837'

# Adjust all table cell font sizes
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('\nComparison: ').bold = True
p.add_run('Outperforms Random Forest (89.90% AUROC), XGBoost (~88%), and Logistic Regression (72.43%). ')
p.add_run('Regression R² of 0.70 indicates 70% variance explained with predictions within ~0.5 log units (MAE).')

p = doc.add_paragraph()
p.add_run('\nKey Innovations: ').bold = True
p.add_run('(1) Selective fine-tuning preserves pretrained chemical knowledge; ')
p.add_run('(2) Cross-modal attention enables dynamic feature weighting; ')
p.add_run('(3) Bilinear fusion captures second-order interactions; ')
p.add_run('(4) Systematic hyperparameter optimization via Optuna TPE.')

# Save document
doc.save('BioActivity_Methodology_Compact.docx')
print("✅ Compact document saved as 'BioActivity_Methodology_Compact.docx'")