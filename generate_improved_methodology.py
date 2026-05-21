#!/usr/bin/env python3
"""
Generate improved methodology DOCX for Q1 journal submission.
Produces: classification/DL/BioActivity_Methodology_Compact.docx  (in-place update)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT_PATH = r'E:\ML\BioActivity\BioActivity_Methodology_Q1Journal.docx'

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))   # 1 cm ≈ 567 twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold_parts=None, italic=False, indent=False):
    """Add a normal paragraph, optionally with inline bold/italic runs."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    if bold_parts is None:
        run = p.add_run(text)
        run.italic = italic
    else:
        # bold_parts: list of (text, is_bold) tuples
        for part_text, is_bold in bold_parts:
            r = p.add_run(part_text)
            r.bold = is_bold
            r.italic = italic
    return p

def add_eq(doc, text):
    """Add a centred equation paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    return p

def table_header_row(table, headers, fill='1F497D', font_color='FFFFFF'):
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        shade_cell(c, fill)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(
            int(font_color[0:2], 16),
            int(font_color[2:4], 16),
            int(font_color[4:6], 16))
        r.font.size = Pt(9.5)

def fill_row(table, row_idx, values, fill=None, bold=False, center=False):
    row = table.rows[row_idx]
    for j, v in enumerate(values):
        c = row.cells[j]
        if fill:
            shade_cell(c, fill)
        p = c.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(v))
        r.bold = bold
        r.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Document title ────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run(
    'FineTunedBERTaECFP: A Multimodal Deep-Learning Framework '
    'for Bioactivity Prediction Using MoLFormer-XL, '
    'ECFP Fingerprints, and Physicochemical Descriptors')
tr.bold      = True
tr.font.size = Pt(14)
tr.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

doc.add_paragraph()  # spacer

# ── Abstract box (shaded) ─────────────────────────────────────────────────────
abs_table = doc.add_table(rows=1, cols=1)
abs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
abs_cell = abs_table.cell(0, 0)
shade_cell(abs_cell, 'EBF5FB')
abs_p = abs_cell.paragraphs[0]
abs_r = abs_p.add_run(
    'Summary.  '
    'We present FineTunedBERTaECFP, a multimodal neural architecture that '
    'integrates three complementary molecular representations for bioactivity '
    'classification: (i) contextual SMILES embeddings from a selectively '
    'fine-tuned MoLFormer-XL transformer (last 12 layers unfrozen), (ii) Extended-Connectivity Fingerprints '
    '(ECFP4/1024-bit), and (iii) a panel of seven physicochemical descriptors. '
    'A cross-modal multi-head attention mechanism (H\u2009=\u20098 heads) dynamically aligns fingerprint and '
    'descriptor features against the SMILES context, and a gated fusion module learns a '
    'sigmoid gate vector that selectively suppresses uninformative embedding dimensions '
    'before classification. Systematic hyperparameter optimisation via Optuna TPE '
    '(50 trials, MedianPruner) yields a configuration attaining '
    'AUROC\u2009=\u200995.39\u202f%, Balanced Accuracy\u2009=\u200992.01\u202f%, and MCC\u2009=\u20090.8281 on a held-out test '
    'set of 952 molecules.')
abs_r.font.size = Pt(9.5)
abs_r.italic    = True
set_col_width(abs_cell, 16.5)

doc.add_paragraph()  # spacer

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 – DATASET AND PREPROCESSING
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Dataset and Molecular Preprocessing', level=1)

add_para(doc,
    'The dataset comprises 6,341 small-molecule compounds labelled as '
    'Active (1) or Inactive (0) towards a defined biological target.  '
    'Each molecule is represented by its canonical SMILES string together '
    'with a set of pre-computed physicochemical properties.')

add_para(doc, 'Data partitioning follows a stratified 70 / 15 / 15 split '
    '(training / validation / test) with a fixed random seed (42) to ensure '
    'reproducibility.  Class proportions (≈19 % active) are preserved in every '
    'split to avoid distribution shift.')

add_para(doc, 'Physicochemical descriptors are computed via RDKit and comprise:',
         bold_parts=[('Physicochemical descriptors ', True),
                     ('are computed via RDKit and comprise:', False)])

for d in [
    'Molecular Weight (MW)',
    'Octanol–Water Partition Coefficient (LogP)',
    'Number of Hydrogen-Bond Donors (NumHDonors)',
    'Topological Polar Surface Area (TPSA)',
    'Number of Rotatable Bonds (NumRotatableBonds)',
    'Carbon sp³ Fraction (FractionCSP3)',
    'Ring Count (RingCount)',
]:
    add_bullet(doc, d)

add_para(doc,
    'All descriptors are standardised using a StandardScaler fitted '
    'exclusively on the training set to prevent data leakage; the same '
    'scaler parameters (μ, σ) are applied to the validation and test sets.')

add_para(doc,
    'SMILES data augmentation is applied during training by stochastically '
    're-ordering atom indices to produce alternative valid SMILES strings '
    'for the same molecule (random SMILES), thereby improving '
    'generalisation without additional data collection.')

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 – MOLECULAR REPRESENTATION ENCODING
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Molecular Representation Encoding', level=1)

add_heading(doc, '2.1  SMILES Encoding via MoLFormer-XL', level=2)

add_para(doc,
    'The SMILES branch uses IBM MoLFormer-XL-both-10pct, a 12-layer '
    'transformer pre-trained on 1.1 billion molecules with rotary positional '
    'embeddings and linear attention (d_model = 768, H = 12 attention heads, '
    'd_ff = 3,072).  To leverage pre-trained chemical knowledge while adapting '
    'to the target task, all parameters are frozen except those of the last '
    'five encoder layers (selective fine-tuning).')

add_para(doc, 'For an input token sequence x = [x₁, x₂, …, xₙ] '
         '(max length 512), the encoder produces:')
add_eq(doc, 'H = MoLFormer(x)  ∈  ℝⁿˣ⁷⁶⁸')

add_para(doc,
    'Three pooling strategies are averaged to form a robust sentence-level '
    'representation that combines positional (CLS), distributional (mean), '
    'and extremal (max) information:')
add_eq(doc,
    'h_SMILES = ( h_CLS + h_mean + h_max ) / 3  ∈  ℝ⁷⁶⁸')
add_para(doc,
    'where  h_CLS = H[:,0,:],  h_mean = (1/n) Σᵢ H[:,i,:],  '
    'h_max = maxᵢ H[:,i,:]  (element-wise maximum).',
    italic=True)

add_heading(doc, '2.2  ECFP Fingerprint Projection', level=2)

add_para(doc,
    'Morgan (ECFP) fingerprints with radius r = 2 and 1,024 bits '
    'are generated through RDKit\'s MorganGenerator, yielding binary '
    'vectors  f_ECFP ∈ {0,1}¹⁰²⁴  that capture circular substructure '
    'environments up to two bond hops from every atom.  A two-layer MLP '
    'projects f_ECFP into the shared 768-dimensional embedding space:')

add_eq(doc, 'h_ECFP = MLP_ECFP( f_ECFP )  ∈  ℝ⁷⁶⁸')
add_para(doc,
    'MLP_ECFP:  Linear(1024→512) → LayerNorm → GELU → Dropout(0.15) → '
    'Linear(512→768) → LayerNorm → GELU → Dropout(0.30)',
    italic=True)

add_heading(doc, '2.3  Physicochemical Descriptor Projection', level=2)

add_para(doc,
    'The seven normalised descriptors  d ∈ ℝ⁷  are mapped to the same '
    'latent space through a deeper three-layer MLP to compensate for their '
    'low dimensionality:')
add_eq(doc, 'h_desc = MLP_desc( d_norm )  ∈  ℝ⁷⁶⁸')
add_para(doc,
    'MLP_desc:  Linear(7→256) → LN → GELU → Dropout(0.15) → '
    'Linear(256→512) → LN → GELU → Dropout(0.15) → '
    'Linear(512→768) → LN → GELU → Dropout(0.30)',
    italic=True)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 – CROSS-MODAL ATTENTION MECHANISM
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  Cross-Modal Multi-Head Attention', level=1)

add_para(doc,
    'After the three encoders, structural (ECFP) and physicochemical (descriptor) '
    'embeddings are stacked along the sequence dimension to form a query tensor, '
    'while the SMILES embedding serves as both key and value.  This design '
    'allows fingerprint and descriptor features to selectively attend to the '
    'parts of the molecular sequence most relevant to their respective '
    'information content:')

add_eq(doc, 'Q  =  stack( [h_ECFP, h_desc], dim=1 )  ∈  ℝᴮˣ²ˣ⁷⁶⁸')
add_eq(doc, 'K  =  V  =  h_SMILES.unsqueeze(1)        ∈  ℝᴮˣ¹ˣ⁷⁶⁸')

add_para(doc,
    'The standard scaled dot-product multi-head attention (H = 8 heads, '
    'd_k = d_model / H = 96) is computed as:')
add_eq(doc, 'headᵢ = softmax( Q Wᵢᴿ (K Wᵢᴷ)ᵀ / √d_k ) · V Wᵢᵛ')
add_eq(doc, 'MultiHead(Q, K, V) = Concat(head₁, …, headₕ) Wᴼ  ∈  ℝᴮˣ²ˣ⁷⁶⁸')
add_para(doc,
    'where  Wᵢᴿ, Wᵢᴷ, Wᵢᵛ ∈ ℝ⁷⁶⁸ˣ⁶⁴  and  Wᴼ ∈ ℝ⁷⁶⁸ˣ⁷⁶⁸  are learned '
    'projection matrices, and the softmax normalises across the single SMILES '
    'token (key–value dimension = 1).',
    italic=True)

add_para(doc,
    'Residual connections with Layer Normalisation prevent representation '
    'collapse and stabilise gradient flow:')
add_eq(doc, 'h_ECFP_attn = LayerNorm( h_ECFP + MultiHead(Q, K, V)[:,0,:] )  ∈  ℝ⁷⁶⁸')
add_eq(doc, 'h_desc_attn = LayerNorm( h_desc + MultiHead(Q, K, V)[:,1,:] )  ∈  ℝ⁷⁶⁸')


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 – GATED FEATURE FUSION
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Gated Feature Fusion', level=1)

add_para(doc,
    'After cross-modal attention, the three representation streams — '
    'h_SMILES, h_ECFP_attn, and h_desc_attn — are integrated through a '
    'gated fusion mechanism.  All three vectors are first concatenated '
    'to form a joint representation  z:')
add_eq(doc, 'z  =  Concat( h_SMILES,  h_ECFP_attn,  h_desc_attn )  ∈  ℝ\u00b2\u00b3\u2070\u2074')

add_para(doc,
    'A learned sigmoid gate vector  g ∈ ℝ⁷⁶⁸  is computed from  z  through '
    'a linear projection followed by a sigmoid non-linearity.  This gate '
    'acts as a per-dimension soft switch, independently amplifying or '
    'suppressing each embedding component based on how informative it is '
    'for the prediction task:')
add_eq(doc, 'g  =  \u03c3( W_g · z + b_g ),    W_g ∈ ℝ⁷⁶⁸ˣ\u00b2\u00b3\u2070\u2074,    g ∈ ℝ⁷⁶⁸')

add_para(doc,
    'A separate linear projection maps  z  into the same 768-dimensional '
    'space, and the final fused representation is produced by an '
    'element-wise (Hadamard) product with the gate:')
add_eq(doc, 'h_fused  =  g  ⊙  ( W_f · z + b_f ),    W_f ∈ ℝ⁷⁶⁸ˣ\u00b2\u00b3\u2070\u2074,    h_fused ∈ ℝ⁷⁶⁸')

add_para(doc,
    'The element-wise product (⊙) ensures that each output dimension of '
    'the fused representation is gated by a value in (0, 1).  Dimensions '
    'encoded by the gate close to 0 are effectively suppressed, while '
    'those close to 1 are passed through unchanged.  This provides a '
    'compact 768-dimensional fused embedding that discards redundant '
    'cross-modal information before the classification head, without '
    'tripling the classifier input size (as would occur with simple '
    'concatenation) or requiring a quadratic weight tensor (as in bilinear fusion).',
    italic=True)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 – CLASSIFICATION HEAD
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Classification Head', level=1)

add_para(doc,
    'The fused representation h_fused ∈ ℝ⁷⁶⁸ is passed through a '
    'progressively contracting MLP classifier with N = 4 layers, '
    'halving the hidden dimension at each step to force increasingly '
    'abstract representations:')

add_eq(doc, 'h⁽⁰⁾  =  h_fused  ∈  ℝ⁷⁶⁸')
add_eq(doc, 'h⁽ⁱ⁾  =  Dropout_{pᵢ}( GELU( LayerNorm( Wᵢ h⁽ⁱ⁻¹⁾ + bᵢ ) ) ),    i = 1, …, 4')
add_eq(doc, 'ŷ      =  w_out\u1D40 h⁽⁴⁾  ∈  ℝ')

add_para(doc, 'Dimensionality trace '
         '(hidden_dim = 768, num_classifier_layers = 4):',
         bold_parts=[('Dimensionality trace ', True),
                     ('(hidden_dim = 768, num_classifier_layers = 4):', False)])
for layer_text in [
    'Layer 1:  ℝ⁷⁶⁸  →  Linear(768→768) + LayerNorm + GELU + Dropout(0.30)',
    'Layer 2:  ℝ⁷⁶⁸  →  Linear(768→384) + LayerNorm + GELU + Dropout(0.30)',
    'Layer 3:  ℝ³⁸⁴  →  Linear(384→192) + LayerNorm + GELU + Dropout(0.30)',
    'Layer 4:  ℝ¹⁹²  →  Linear(192→96)  + LayerNorm + GELU + Dropout(0.15)',
    'Output:   ℝ⁹⁶   →  Linear(96→1)    =  ŷ',
]:
    add_bullet(doc, layer_text)

add_para(doc,
    'Dropout probability is halved in the final layer to retain predictive '
    'signal close to the output.  Batch normalisation was replaced by '
    'LayerNorm to avoid dependence on batch statistics at inference time.')


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 – TRAINING PROCEDURE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Training Procedure', level=1)

add_heading(doc, '6.1  Loss Function', level=2)
add_para(doc,
    'Class imbalance (≈79 % inactive vs 21 % active) is addressed through '
    'a class-weighted Binary Cross-Entropy (BCE) loss with a per-sample '
    'positive weight  w_pos = N_neg / N_pos:')
add_eq(doc,
    'ℒ(ŷ, y) = − (1/B) Σᵢ [ w_pos · yᵢ log σ(ŷᵢ) + (1−yᵢ) log(1−σ(ŷᵢ)) ]')
add_para(doc,
    'where  B  is the mini-batch size,  σ(·)  is the sigmoid function, '
    'and  w_pos  penalises false negatives proportionally to the '
    'imbalance ratio.',
    italic=True)

add_heading(doc, '6.2  Optimiser and Learning-Rate Schedule', level=2)
add_para(doc,
    'RMSprop (momentum = 0.9, ε = 10⁻⁸) is used as the base optimiser '
    'with weight decay  λ = 1.24×10⁻⁴.  The learning rate follows a '
    'One-Cycle cosine schedule:')
add_eq(doc,
    'α(t) = αₘₐₓ · [ (1 − t/T) · cos(π·t/T) + 1 ] / 2  +  αₘᵢₙ')
add_para(doc,
    'where  αₘₐₓ = 2.37×10⁻⁶  is the peak learning rate, '
    'αₘᵢₙ = αₘₐₓ / 25,  t  is the current step, and '
    'T  is the total number of training steps.  '
    'A warm-up phase covering 10 % of T brings the learning rate from '
    'αₘᵢₙ to αₘₐₓ before the cosine annealing begins.',
    italic=True)

add_heading(doc, '6.3  Regularisation and Training Details', level=2)
for item in [
    'Dropout: p = 0.30 in all intermediate layers; p = 0.15 in final classifier layer and attention.',
    'Gradient clipping: max_norm = 1.0 (L₂-norm of all parameters).',
    'Gradient accumulation: 2 steps → effective batch size = 32 × 2 = 64.',
    'Early stopping: patience = 7 epochs on validation AUROC.',
    'Maximum epochs: 50.  Training hardware: NVIDIA GPU (CUDA).',
    'All random seeds set to 42 (PyTorch, NumPy, Python random).',
]:
    add_bullet(doc, item)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 – HYPERPARAMETER OPTIMISATION
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7.  Hyperparameter Optimisation via Optuna', level=1)

add_para(doc,
    'All architectural and training hyperparameters are jointly optimised '
    'using Optuna\'s Tree-structured Parzen Estimator (TPE) sampler '
    '(n_startup_trials = 5) over 50 trials with a maximum wall-clock '
    'budget of 6 hours.  Poorly performing trials are pruned early by '
    'the MedianPruner (n_warmup_steps = 5 epochs), which terminates a '
    'trial if its intermediate validation AUROC falls below the median of '
    'completed trials at the same epoch.  The objective function is '
    'validation AUROC (maximised).')

# Hyperparameter table
add_para(doc, 'Table 1.  Hyperparameter search space.',
         bold_parts=[('Table 1.  ', True),
                     ('Hyperparameter search space.', False)])
hp_table = doc.add_table(rows=10, cols=4)
hp_table.style = 'Table Grid'
hp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

table_header_row(hp_table,
    ['Hyperparameter', 'Search Range / Options',
     'Hyperparameter', 'Search Range / Options'])

rows_data = [
    ('unfreeze_layers',       'ℤ ∈ [4, 12]',
     'learning_rate',         'Log-uniform ∈ [5×10⁻⁷, 1×10⁻⁵]'),
    ('num_heads',             '{4, 8, 12, 16}',
     'weight_decay',          'Log-uniform ∈ [10⁻⁵, 10⁻³]'),
    ('hidden_dim',            '{256, 512, 768}',
     'batch_size',            '{16, 32, 64}'),
    ('num_classifier_layers', 'ℤ ∈ [3, 5]',
     'grad_accum_steps',      '{1, 2, 4, 8}'),
    ('dropout',               'Uniform ∈ [0.1, 0.6]',
     'optimizer',             '{Adam, AdamW, RMSprop, RAdam}'),
    ('fusion_type',           '{concat, gated, bilinear}',
     'scheduler',             '{CosineAnnealingLR, ReduceLROnPlateau,'),
    ('',                      '',
     '',                       '  OneCycleLR, CosineAnnWarmRestarts, Lambda}'),
]

for i, (a, b, c, d) in enumerate(rows_data):
    fill = 'EBF5FB' if i % 2 == 0 else 'FDFEFE'
    fill_row(hp_table, i + 1, [a, b, c, d], fill=fill if i < 5 else None)
    if i == 5:
        fill_row(hp_table, 6, [a, b, c, d], fill='EBF5FB')
    if i == 6:
        fill_row(hp_table, 7, [a, b, c, d], fill='FDFEFE')

# Best params row
fill_row(hp_table, 8, ['Best Configuration', '', '', ''], bold=True,
         fill='D6EAF8')
fill_row(hp_table, 9,
    ['unfreeze=12, heads=8, hidden=768, dropout=0.30, layers=4, fusion=gated',
     '', '', ''])

doc.add_paragraph()  # spacer

add_para(doc,
    'Best hyperparameters (production model):  unfreeze_layers = 12,  num_heads = 8,  '
    'hidden_dim = 768,  dropout = 0.30,  num_classifier_layers = 4,  '
    'fusion_type = gated,  batch_size = 32,  '
    'learning_rate = 1\u00d710\u207b\u2075,  weight_decay = 0.01,  '
    'optimizer = AdamW,  scheduler = LambdaLR (cosine warm-up).')


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 – RESULTS AND COMPARISON
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8.  Results and Comparison', level=1)

add_para(doc,
    'Table 2 summarises performance on the held-out test set (n = 952).  '
    'The decision threshold is fixed at 0.5.  The Brier score and '
    'Matthews Correlation Coefficient (MCC) provide calibration and '
    'balanced-accuracy assessments, respectively.')

# Results table
add_para(doc, 'Table 2.  Test-set performance of FineTunedBERTaECFP.',
         bold_parts=[('Table 2.  ', True),
                     ('Test-set performance of FineTunedBERTaECFP.', False)])
metrics = [
    ('AUROC',            '95.39 %'),
    ('Balanced Accuracy','92.01 %'),
    ('Accuracy',         '96.90 %'),
    ('Precision',        '97.12 %'),
    ('Recall / Sensitivity', '96.24 %'),
    ('Specificity',      '87.79 %'),
    ('F1 Score',         '96.68 %'),
    ('MCC',              '0.8281'),
    ('AUPRC',            '—'),
    ('Brier Score',      '—'),
    ('NPV',              '—'),
    ('DOR',              '—'),
]
res_table = doc.add_table(rows=len(metrics) + 1, cols=2)
res_table.style = 'Table Grid'
res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(res_table, ['Metric', 'Score'])
for i, (metric, score) in enumerate(metrics):
    fill = 'EBF5FB' if i % 2 == 0 else 'FDFEFE'
    fill_row(res_table, i + 1, [metric, score], fill=fill, center=True)

doc.add_paragraph()

# Additional regression results
add_para(doc, 'Table 3.  Regression sub-task results (pIC₅₀ prediction).',
         bold_parts=[('Table 3.  ', True),
                     ('Regression sub-task results (pIC₅₀ prediction).', False)])
reg_metrics = [
    ('R² (Coefficient of Determination)', '0.7006'),
    ('RMSE (Root Mean Square Error)',      '0.7055  log units'),
    ('MAE (Mean Absolute Error)',          '0.4932  log units'),
    ('Pearson Correlation (r)',            '0.837'),
]
reg_table = doc.add_table(rows=len(reg_metrics) + 1, cols=2)
reg_table.style = 'Table Grid'
reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(reg_table, ['Metric', 'Score'])
for i, (m, s) in enumerate(reg_metrics):
    fill = 'EBF5FB' if i % 2 == 0 else 'FDFEFE'
    fill_row(reg_table, i + 1, [m, s], fill=fill, center=True)

doc.add_paragraph()

# Comparison table
add_para(doc, 'Table 4.  Comparison against baseline classifiers.',
         bold_parts=[('Table 4.  ', True),
                     ('Comparison against baseline classifiers.', False)])
comp = [
    ('FineTunedBERTaECFP (proposed)', '95.39 %', '92.01 %', '0.8281'),
    ('ChemBERTa-zinc (fine-tuned)',   '~93 %',    '~89 %',   '—'),
    ('Random Forest',                 '89.90 %',  '—',       '—'),
    ('XGBoost',                       '~88 %',    '—',       '—'),
    ('Logistic Regression',           '72.43 %',  '—',       '—'),
]
cmp_table = doc.add_table(rows=len(comp) + 1, cols=4)
cmp_table.style = 'Table Grid'
cmp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(cmp_table, ['Model', 'AUROC', 'Balanced Acc.', 'MCC'])
for i, row_data in enumerate(comp):
    fill = 'D5F5E3' if i == 0 else ('EBF5FB' if i % 2 == 0 else 'FDFEFE')
    bold = (i == 0)
    fill_row(cmp_table, i + 1, list(row_data), fill=fill, bold=bold, center=True)


# ═════════════════════════════════════════════════════════════════════════════
# SECTION 9 – KEY INNOVATIONS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9.  Key Methodological Contributions', level=1)

for item in [
    '(1)  Heterogeneous multimodal fusion:  SMILES sequence, binary fingerprint, '
    'and real-valued descriptor streams are projected into a shared latent space, '
    'enabling complementary information to be jointly exploited.',

    '(2)  Selective fine-tuning:  Only the last five transformer encoder layers of '
    'MoLFormer-XL are updated during training, striking a balance between '
    'task specialisation and preservation of broadly applicable chemical '
    'representations learned during pre-training.',

    '(3)  Cross-modal attention:  By treating fingerprint and descriptor embeddings '
    'as queries and the SMILES representation as key–value memory, the network '
    'dynamically weights which structural and physicochemical features are most '
    'informative for each molecule.',

    '(4)  Gated fusion:  A learned sigmoid gate vector  g = σ(W_g·z) ∈ ℝ⁷⁶⁸  '
    'selectively suppresses uninformative dimensions of the concatenated '
    'multimodal representation  z = [h_SMILES; h_ECFP_attn; h_desc_attn],  '
    'yielding a compact 768-dimensional fused embedding without tripling the '
    'classifier input dimension.',

    '(5)  Joint hyperparameter search:  Simultaneously optimising architecture '
    '(fusion type, attention heads, depth) and training (optimiser, schedule, '
    'regularisation) via Optuna TPE ensures that reported performance is not '
    'an artefact of manual tuning.',
]:
    add_para(doc, item)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 10 – IMPLEMENTATION DETAILS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10.  Implementation Details', level=1)

for item in [
    'Framework:  PyTorch ≥ 2.0,  Hugging Face transformers,  RDKit 2024,  '
    'Optuna 3.x,  scikit-learn.',
    'Pre-trained model:  ibm/MoLFormer-XL-both-10pct (HuggingFace Hub).',
    'Tokeniser:  ibm/MoLFormer-XL-both-10pct,  max_length = 512,  truncation = True,  '
    'padding = True.',
    'Hardware:  NVIDIA GPU (CUDA); all experiments reproducible on CPU.',
    'Random seeds:  Python random = 42,  NumPy = 42,  torch.manual_seed = 42.',
    'Code availability:  Provided as a Jupyter notebook '
    '(bioactivity_dl 8_2 C.ipynb) with the supplementary materials.',
]:
    add_bullet(doc, item)


# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f'Saved → {OUT_PATH}')
