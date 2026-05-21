from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title
title = doc.add_heading('Mathematical Formulation of FineTunedBERTaECFP', 0)
for run in title.runs:
    run.font.size = Pt(16)
    run.font.bold = True

# Introduction
p = doc.add_paragraph(
    'This document presents the complete mathematical formulation of the FineTunedBERTaECFP model, '
    'a multimodal deep learning architecture for molecular bioactivity prediction that integrates SMILES sequences, '
    'ECFP fingerprints, and physicochemical descriptors through cross-modal attention and adaptive fusion strategies.'
)

# Section 1
doc.add_heading('1. Selective Transformer Fine-Tuning', 1)
p = doc.add_paragraph(
    'Given a SMILES sequence tokenized as x = [x₁, x₂, ..., xₙ], the pretrained MoLFormer transformer produces '
    'contextualized embeddings:'
)

eq = doc.add_paragraph('H = MoLFormer(x) ∈ ℝⁿˣ⁷⁶⁸', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    'where only the last L ∈ [4, 12] transformer layers have trainable parameters:'
)

eq = doc.add_paragraph('θ_trainable = {θ_layer₁₂₋ₗ₊₁, ..., θ_layer₁₂}', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 2
doc.add_heading('2. SMILES Representation via Triple Pooling', 1)
p = doc.add_paragraph(
    'To obtain a fixed-size molecular representation that captures both global and local information, we employ '
    'a triple-pooling aggregation strategy:'
)

eq = doc.add_paragraph('h_SMILES = (1/3)(h_CLS + h_mean + h_max)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('h_CLS = H[0]  (global representation from [CLS] token)', style='List Bullet')
doc.add_paragraph('h_mean = (1/n)Σⁿᵢ₌₁H[i]  (average pooling across all tokens)', style='List Bullet')
doc.add_paragraph('h_max = maxⁿᵢ₌₁H[i]  (element-wise max pooling)', style='List Bullet')

# Section 3
doc.add_heading('3. ECFP Feature Projection', 1)
p = doc.add_paragraph(
    'For Morgan fingerprint f_ECFP ∈ {0,1}¹⁰²⁴, we apply a two-layer projection network with progressive '
    'dimensionality reduction:'
)

eq = doc.add_paragraph('h_ECFP⁽¹⁾ = Dropout₀.₅ₚ(GELU(LN(W₁ᴱᶜᶠᴾ · f_ECFP + b₁ᴱᶜᶠᴾ)))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_ECFP = Dropoutₚ(GELU(LN(W₂ᴱᶜᶠᴾ · h_ECFP⁽¹⁾ + b₂ᴱᶜᶠᴾ)))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('W₁ᴱᶜᶠᴾ ∈ ℝ⁵¹²ˣ¹⁰²⁴  (first projection matrix)', style='List Bullet')
doc.add_paragraph('W₂ᴱᶜᶠᴾ ∈ ℝ⁷⁶⁸ˣ⁵¹²  (second projection matrix)', style='List Bullet')
doc.add_paragraph('h_ECFP ∈ ℝ⁷⁶⁸  (final ECFP embedding)', style='List Bullet')
doc.add_paragraph('LN denotes Layer Normalization', style='List Bullet')
doc.add_paragraph('GELU is the Gaussian Error Linear Unit activation function', style='List Bullet')

# Section 4
doc.add_heading('4. Molecular Descriptor Projection', 1)
p = doc.add_paragraph(
    'For descriptor vector d ∈ ℝ⁷ containing MW, LogP, NumHDonors, TPSA, NumRotatableBonds, FractionCSP3, '
    'and RingCount, we apply a three-layer MLP with progressive expansion:'
)

eq = doc.add_paragraph('h_desc⁽¹⁾ = Dropout₀.₅ₚ(GELU(LN(W₁ᵈᵉˢᶜ · d + b₁ᵈᵉˢᶜ)))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_desc⁽²⁾ = Dropout₀.₅ₚ(GELU(LN(W₂ᵈᵉˢᶜ · h_desc⁽¹⁾ + b₂ᵈᵉˢᶜ)))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_desc = Dropoutₚ(GELU(LN(W₃ᵈᵉˢᶜ · h_desc⁽²⁾ + b₃ᵈᵉˢᶜ)))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where dimensions progress as: 7 → 256 → 512 → 768.')

# Section 5
doc.add_heading('5. Cross-Modal Multi-Head Attention', 1)
p = doc.add_paragraph(
    'The query matrix Q comprises stacked ECFP and descriptor embeddings, while SMILES serves as key and value:'
)

eq = doc.add_paragraph('Q = [h_ECFP; h_desc] ∈ ℝ²ˣ⁷⁶⁸,    K = V = h_SMILES ∈ ℝ¹ˣ⁷⁶⁸', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('Multi-head attention with H heads is computed as:')

eq = doc.add_paragraph('MultiHead(Q, K, V) = Concat(head₁, ..., headₕ)W^O', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('headᵢ = Attention(QWᵢ^Q, KWᵢ^K, VWᵢ^V)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('Attention(Q\', K\', V\') = softmax((Q\'K\'^T)/√dₖ)V\'', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('Wᵢ^Q, Wᵢ^K, Wᵢ^V ∈ ℝ⁷⁶⁸ˣᵈᵏ with dₖ = 768/H', style='List Bullet')
doc.add_paragraph('W^O ∈ ℝ⁷⁶⁸ˣ⁷⁶⁸ is the output projection matrix', style='List Bullet')

p = doc.add_paragraph('The attention output is combined with residual connections and layer normalization:')

eq = doc.add_paragraph('h_ECFP^attn = LN(h_ECFP + MultiHead(h_ECFP, h_SMILES, h_SMILES)[0])', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_desc^attn = LN(h_desc + MultiHead(h_desc, h_SMILES, h_SMILES)[1])', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 6
doc.add_heading('6. Multimodal Fusion Strategies', 1)
p = doc.add_paragraph(
    'After obtaining attention-refined representations, we combine them using one of three fusion strategies:'
)

doc.add_heading('6.1 Concatenation Fusion', 2)
p = doc.add_paragraph('Direct concatenation preserving all information:')

eq = doc.add_paragraph('h_fused = [h_SMILES; h_ECFP^attn; h_desc^attn] ∈ ℝ²³⁰⁴', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('6.2 Gated Fusion', 2)
p = doc.add_paragraph('Learnable gates for adaptive modality weighting:')

eq = doc.add_paragraph('z = [h_SMILES; h_ECFP^attn; h_desc^attn] ∈ ℝ²³⁰⁴', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('g = σ(W_gate · z + b_gate) ∈ ℝ⁷⁶⁸', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_fused = g ⊙ (W_proj · z + b_proj)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where σ is the sigmoid activation and ⊙ denotes element-wise multiplication (Hadamard product).')

doc.add_heading('6.3 Bilinear Fusion (Best Performing)', 2)
p = doc.add_paragraph('Captures second-order multiplicative interactions between modalities:')

eq = doc.add_paragraph('h_struct = h_ECFP^attn + h_desc^attn', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h_fused = h_SMILES^T W_bilinear h_struct', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where W_bilinear ∈ ℝ⁷⁶⁸ˣ⁷⁶⁸ is a learnable bilinear transformation matrix.')

# Section 7
doc.add_heading('7. Classification Head with Progressive Reduction', 1)
p = doc.add_paragraph('For N classification layers with progressive halving of dimensions:')

eq = doc.add_paragraph('h⁽⁰⁾ = h_fused', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('h⁽ⁱ⁾ = Dropoutₚᵢ(GELU(LN(Wᵢh⁽ⁱ⁻¹⁾ + bᵢ))),    i = 1, ..., N', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('Dimension dᵢ = d_hidden/2ⁱ (progressive halving)', style='List Bullet')
doc.add_paragraph('Dropout rate pᵢ = p for i < N, else pᵢ = 0.5p (reduced in final layer)', style='List Bullet')

p = doc.add_paragraph('Final prediction:')

eq = doc.add_paragraph('ŷ = w_out^T h⁽ᴺ⁾ + b_out', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('p(active) = σ(ŷ) = 1/(1 + e^(-ŷ))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 8
doc.add_heading('8. Loss Function', 1)
p = doc.add_paragraph('Weighted binary cross-entropy to handle class imbalance:')

eq = doc.add_paragraph('ℒ = -(1/B)Σᵢ₌₁ᴮ[w_pos · yᵢlog(σ(ŷᵢ)) + (1-yᵢ)log(1-σ(ŷᵢ))]', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('B is the batch size', style='List Bullet')
doc.add_paragraph('yᵢ ∈ {0, 1} is the true binary label', style='List Bullet')
doc.add_paragraph('w_pos = N_neg/N_pos is the positive class weight (inverse frequency weighting)', style='List Bullet')
doc.add_paragraph('σ(·) is the sigmoid function', style='List Bullet')

# Section 9
doc.add_heading('9. Optimization Strategy', 1)

doc.add_heading('9.1 RMSprop Optimizer', 2)
p = doc.add_paragraph('Root Mean Square Propagation with momentum update:')

eq = doc.add_paragraph('v_t = β · v_(t-1) + (1-β) · g_t²', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

eq = doc.add_paragraph('θ_t = θ_(t-1) - α/√(v_t + ε) · g_t', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where β = 0.9, ε = 10⁻⁸, and α is the learning rate.')

doc.add_heading('9.2 OneCycleLR Scheduler', 2)
p = doc.add_paragraph('Cyclical learning rate with cosine annealing:')

eq = doc.add_paragraph('α(t) = α_max · [(1 - t/T)/(cos(π·t/T) + 1)] · 0.5 + α_min', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('t is the current training step', style='List Bullet')
doc.add_paragraph('T is the total number of training steps', style='List Bullet')
doc.add_paragraph('α_max is the maximum learning rate', style='List Bullet')
doc.add_paragraph('α_min = α_max/25 is the minimum learning rate', style='List Bullet')

doc.add_heading('9.3 Gradient Clipping', 2)
p = doc.add_paragraph('To prevent exploding gradients during backpropagation:')

eq = doc.add_paragraph('g ← g · min(1, τ/||g||₂)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where τ = 1.0 is the maximum gradient norm (L2-norm clipping).')

# Section 10
doc.add_heading('10. Hyperparameter Optimization', 1)
p = doc.add_paragraph(
    'We employ Optuna with Tree-structured Parzen Estimator (TPE) sampling for Bayesian hyperparameter optimization. '
    'The TPE algorithm models the objective function using two probability distributions:'
)

eq = doc.add_paragraph('p(x|y) = { ℓ(x) if y < y*, g(x) if y ≥ y* }', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where:')
doc.add_paragraph('y* is a performance threshold (e.g., median AUROC)', style='List Bullet')
doc.add_paragraph('ℓ(x) models hyperparameters that led to good performance', style='List Bullet')
doc.add_paragraph('g(x) models hyperparameters that led to poor performance', style='List Bullet')

p = doc.add_paragraph('MedianPruner terminates unpromising trials early. A trial is pruned at epoch e if:')

eq = doc.add_paragraph('AUROC_e < median{AUROC_e,1, AUROC_e,2, ..., AUROC_e,k}', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('where k is the number of completed trials at epoch e.')

# Section 11
doc.add_heading('11. Performance Metrics', 1)

p = doc.add_paragraph()
p.add_run('Area Under ROC Curve (AUROC):').bold = True

eq = doc.add_paragraph('AUROC = ∫₀¹ TPR(t) d(FPR(t))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Balanced Accuracy:').bold = True

eq = doc.add_paragraph('Balanced Acc = (TPR + TNR)/2', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Matthews Correlation Coefficient:').bold = True

eq = doc.add_paragraph('MCC = (TP·TN - FP·FN) / √((TP+FP)(TP+FN)(TN+FP)(TN+FN))', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('F1 Score:').bold = True

eq = doc.add_paragraph('F1 = 2 · (Precision · Recall)/(Precision + Recall)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Precision and Recall:').bold = True

eq = doc.add_paragraph('Precision = TP/(TP + FP),    Recall = TP/(TP + FN)', style='Normal')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 12
doc.add_heading('12. Results Summary', 1)

p = doc.add_paragraph()
p.add_run('Best Hyperparameters:').bold = True
doc.add_paragraph('Unfrozen layers: L = 5', style='List Bullet')
doc.add_paragraph('Attention heads: H = 12', style='List Bullet')
doc.add_paragraph('Fusion strategy: Bilinear', style='List Bullet')
doc.add_paragraph('Dropout rate: p = 0.30', style='List Bullet')
doc.add_paragraph('Batch size: 32 with 2-step gradient accumulation', style='List Bullet')
doc.add_paragraph('Optimizer: RMSprop with learning rate α = 3.2×10⁻⁶', style='List Bullet')
doc.add_paragraph('Scheduler: OneCycleLR', style='List Bullet')

p = doc.add_paragraph()
p.add_run('\nClassification Performance:').bold = True
doc.add_paragraph('AUROC: 95.39%', style='List Bullet')
doc.add_paragraph('Balanced Accuracy: 92.01%', style='List Bullet')
doc.add_paragraph('Precision: 97.12%', style='List Bullet')
doc.add_paragraph('Recall: 96.24%', style='List Bullet')
doc.add_paragraph('F1-Score: 94.60%', style='List Bullet')
doc.add_paragraph('Matthews Correlation Coefficient: 0.8281', style='List Bullet')

p = doc.add_paragraph()
p.add_run('\nRegression Performance:').bold = True
doc.add_paragraph('R² Score: 0.7006 (70.06% variance explained)', style='List Bullet')
doc.add_paragraph('Root Mean Squared Error (RMSE): 0.7055', style='List Bullet')
doc.add_paragraph('Mean Absolute Error (MAE): 0.4932', style='List Bullet')
doc.add_paragraph('Pearson Correlation Coefficient: 0.837', style='List Bullet')

# Conclusion
doc.add_heading('13. Conclusion', 1)
p = doc.add_paragraph(
    'The FineTunedBERTaECFP model achieves state-of-the-art performance on molecular bioactivity prediction '
    'through: (1) selective fine-tuning of pretrained molecular transformers, (2) cross-modal attention enabling '
    'dynamic feature weighting, (3) bilinear fusion capturing second-order interactions, and (4) systematic '
    'hyperparameter optimization. The model significantly outperforms classical baselines (Random Forest: 89.90% AUROC, '
    'XGBoost: ~88% AUROC) with 95.39% AUROC for classification and R² = 0.70 for regression, demonstrating robust '
    'predictive capability across both discrete and continuous bioactivity prediction tasks.'
)

# Save document
doc.save('BioActivity_Mathematical_Formulation.docx')
print("✅ Document saved as 'BioActivity_Mathematical_Formulation.docx'")
