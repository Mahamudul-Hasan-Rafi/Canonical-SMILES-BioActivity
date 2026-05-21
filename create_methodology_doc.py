from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set narrow margins for compactness
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title
title = doc.add_heading('Methodology: Multimodal Bioactivity Prediction with IBM MoLFormer-XL', 0)
title_format = title.paragraph_format
title_format.space_after = Pt(6)
for run in title.runs:
    run.font.size = Pt(14)

# Overview - compact
p = doc.add_paragraph(
    'The FineTunedBERTaECFP model integrates three molecular representations: SMILES (via IBM MoLFormer-XL-both-10pct), '
    'Extended-Connectivity Fingerprints (ECFP), and physicochemical descriptors through cross-modal attention and fusion mechanisms.'
)
p.paragraph_format.space_after = Pt(4)

# 1. Model Architecture
h1 = doc.add_heading('1. Model Architecture', 1)
for run in h1.runs:
    run.font.size = Pt(12)
h1.paragraph_format.space_before = Pt(6)
h1.paragraph_format.space_after = Pt(4)

# 1.1 SMILES Encoding
h2 = doc.add_heading('1.1 SMILES Encoding', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph(
    'Pre-trained MoLFormer-XL transformer with selective fine-tuning of last L layers (L ∈ [4,12]). '
    'For tokenized SMILES x = [x₁, x₂, ..., xₙ]:'
)
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('H = MoLFormer(x) ∈ ℝⁿˣ⁷⁶⁸')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('Triple-pooling aggregation:')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h_SMILES = (1/3)(h_CLS + h_mean + h_max)')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(4)

# 1.2 ECFP Projection
h2 = doc.add_heading('1.2 ECFP Projection', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('Morgan fingerprints (r=2, b=1024) projected via two-layer MLP:')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h_ECFP = Dropout(GELU(LayerNorm(W₂(GELU(LayerNorm(W₁f_ECFP))))))')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('W₁ ∈ ℝ⁵¹²ˣ¹⁰²⁴, W₂ ∈ ℝ⁷⁶⁸ˣ⁵¹², h_ECFP ∈ ℝ⁷⁶⁸')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(4)

# 1.3 Descriptor Projection
h2 = doc.add_heading('1.3 Molecular Descriptors', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('Seven standardized descriptors (MW, LogP, NumHDonors, TPSA, NumRotatableBonds, FractionCSP3, RingCount) via 3-layer MLP (7→256→512→768):')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h_desc = MLP₃(d), d ∈ ℝ⁷, h_desc ∈ ℝ⁷⁶⁸')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(4)

# 1.4 Cross-Modal Attention
h2 = doc.add_heading('1.4 Cross-Modal Attention', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('Multi-head attention (H=12 heads) with SMILES as key-value, ECFP/descriptors as queries:')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('Attention(Q, K, V) = softmax(QKᵀ/√dₖ)V')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('Q = [h_ECFP; h_desc] ∈ ℝ²ˣ⁷⁶⁸,  K = V = h_SMILES ∈ ℝ¹ˣ⁷⁶⁸')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('With residual connections and layer normalization:')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h_ECFP^attn = LayerNorm(h_ECFP + Attention(h_ECFP, h_SMILES, h_SMILES))')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h_desc^attn = LayerNorm(h_desc + Attention(h_desc, h_SMILES, h_SMILES))')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(4)

# 1.5 Fusion
h2 = doc.add_heading('1.5 Multimodal Fusion', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('Three fusion strategies (selected via hyperparameter optimization):')
p.paragraph_format.space_after = Pt(2)

# Create compact fusion table
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.autofit = True

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Strategy'
hdr_cells[1].text = 'Formulation'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

strategies = [
    ('Concatenation', 'h_fused = [h_SMILES; h_ECFP^attn; h_desc^attn] ∈ ℝ²³⁰⁴'),
    ('Gated', 'h_fused = σ(Wₘz + bₘ) ⊙ (Wfz + bf) ∈ ℝ⁷⁶⁸'),
    ('Bilinear (Best)', 'h_fused = h_SMILESᵀ Wᵦ (h_ECFP^attn + h_desc^attn) ∈ ℝ⁷⁶⁸')
]

for i, (strategy, formula) in enumerate(strategies, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = strategy
    row_cells[1].text = formula
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# 1.6 Classification Head
h2 = doc.add_heading('1.6 Classification Head', 2)
for run in h2.runs:
    run.font.size = Pt(11)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('N-layer MLP (N=4) with progressive dimension halving:')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('h⁽ⁱ⁾ = Dropout(GELU(LayerNorm(Wᵢh⁽ⁱ⁻¹⁾))), i = 1,...,N')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('ŷ = Wₒᵤₜh⁽ᴺ⁾ + bₒᵤₜ,  p(active) = σ(ŷ)')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(4)

# 2. Training
doc.add_page_break()

h1 = doc.add_heading('2. Training Configuration', 1)
for run in h1.runs:
    run.font.size = Pt(12)
h1.paragraph_format.space_before = Pt(6)
h1.paragraph_format.space_after = Pt(4)

# Loss
p = doc.add_paragraph()
p.add_run('Loss Function: ').bold = True
p.add_run('Weighted binary cross-entropy for class imbalance (w_pos = N_neg/N_pos):')
p.paragraph_format.space_after = Pt(2)

eq = doc.add_paragraph('ℒ = -(1/B)Σᵢ[wₚₒₛ·yᵢlog(σ(ŷᵢ)) + (1-yᵢ)log(1-σ(ŷᵢ))]')
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eq.runs:
    run.font.size = Pt(10)
eq.paragraph_format.space_before = Pt(2)
eq.paragraph_format.space_after = Pt(4)

# Optimizer
p = doc.add_paragraph()
p.add_run('Optimizer: ').bold = True
p.add_run('RMSprop with learning rate α ∈ [5×10⁻⁷, 1×10⁻⁵], weight decay λ ∈ [1×10⁻⁵, 1×10⁻³]')
p.paragraph_format.space_after = Pt(4)

# Regularization
p = doc.add_paragraph()
p.add_run('Regularization: ').bold = True
p.add_run('Dropout p ∈ [0.1, 0.6], gradient clipping (max norm = 1.0), layer normalization')
p.paragraph_format.space_after = Pt(4)

# Scheduler
p = doc.add_paragraph()
p.add_run('LR Scheduler: ').bold = True
p.add_run('OneCycleLR with cosine annealing')
p.paragraph_format.space_after = Pt(4)

# Hyperparameter Optimization
p = doc.add_paragraph()
p.add_run('Hyperparameter Optimization: ').bold = True
p.add_run('Optuna TPE sampler, 50 trials, MedianPruner for early stopping')
p.paragraph_format.space_after = Pt(6)

# 3. Results
h1 = doc.add_heading('3. Performance Results', 1)
for run in h1.runs:
    run.font.size = Pt(12)
h1.paragraph_format.space_before = Pt(6)
h1.paragraph_format.space_after = Pt(4)

# Best Configuration
p = doc.add_paragraph()
p.add_run('Best Configuration: ').bold = True
p.add_run('L=5 unfrozen layers, H=12 heads, bilinear fusion, batch 32, RMSprop + OneCycleLR')
p.paragraph_format.space_after = Pt(4)

# Results table
table = doc.add_table(rows=3, cols=7)
table.style = 'Light Grid Accent 1'

# Header
hdr_cells = table.rows[0].cells
headers = ['Task', 'AUROC', 'Balanced Acc.', 'Precision', 'Recall', 'F1-Score', 'MCC/R²']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

# Classification results
row1_cells = table.rows[1].cells
class_results = ['Classification', '95.39%', '92.01%', '97.12%', '96.24%', '94.60%', '0.8281']
for i, result in enumerate(class_results):
    row1_cells[i].text = result
    for paragraph in row1_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Regression results
row2_cells = table.rows[2].cells
reg_results = ['Regression', '—', '—', '—', '—', 'RMSE: 0.7055', 'R²: 0.7006']
for i, result in enumerate(reg_results):
    row2_cells[i].text = result
    for paragraph in row2_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)

# Key hyperparameters table
p = doc.add_paragraph()
p.add_run('Key Hyperparameters:').bold = True
p.paragraph_format.space_after = Pt(2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

params = [
    ('unfreeze_layers', '[4,12]', 'num_heads', '{4,8,12,16}'),
    ('dropout', '[0.1,0.6]', 'fusion_type', '{concat,gated,bilinear}'),
    ('batch_size', '{16,32,64}', 'learning_rate', '[5×10⁻⁷,1×10⁻⁵]'),
    ('num_layers', '[3,5]', 'optimizer', '{Adam,RMSprop,RAdam}')
]

for i, (p1, r1, p2, r2) in enumerate(params):
    row_cells = table.rows[i].cells
    row_cells[0].text = p1
    row_cells[1].text = r1
    row_cells[2].text = p2
    row_cells[3].text = r2
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.add_run('Dataset: ').italic = True
p.add_run('6,341 molecular compounds with SMILES, ECFP fingerprints, and 7 physicochemical descriptors. Train/Val/Test split: 70/15/15.')
for run in p.runs:
    run.font.size = Pt(9)

# Save document
doc.save('BioActivity_Methodology_Compact.docx')
print("✅ Compact methodology document created: BioActivity_Methodology_Compact.docx")
